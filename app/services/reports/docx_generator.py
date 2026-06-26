"""DOCX report generator menggunakan python-docx."""

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.util import Cm
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    Document = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    qn = None  # type: ignore
    Inches = Pt = RGBColor = Cm = None  # type: ignore

from app.services.reports.schemas import ReportData

# Warna tema — hanya dibuat jika python-docx tersedia
if _DOCX_AVAILABLE:
    COLOR_PRIMARY = RGBColor(0x1e, 0x3a, 0x5f)
    COLOR_POSITIVE = RGBColor(0x27, 0xae, 0x60)
    COLOR_NEGATIVE = RGBColor(0xe7, 0x4c, 0x3c)
    COLOR_NEUTRAL = RGBColor(0x95, 0xa5, 0xa6)
    COLOR_ACCENT = RGBColor(0x2e, 0x86, 0xde)
else:
    COLOR_PRIMARY = COLOR_POSITIVE = COLOR_NEGATIVE = COLOR_NEUTRAL = COLOR_ACCENT = None


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set background warna cell DOCX."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = cell._tc.new_element("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


class DOCXReportGenerator:
    def generate(self, data: ReportData, output_dir: str) -> str:
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        file_path = str(Path(output_dir) / f"{data.report_id}.docx")

        doc = Document()

        # ── Margin ──────────────────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # ── Cover ───────────────────────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(data.title)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_PRIMARY

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Keyword: {data.keyword_text}  |  "
            f"Dibuat: {data.generated_at.strftime('%d %b %Y %H:%M')} UTC  |  "
            f"Periode: {data.period}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        # ── Ringkasan Data ──────────────────────────────────────────────────────
        doc.add_heading("Ringkasan Data", level=1)

        stat_table = doc.add_table(rows=2, cols=3)
        stat_table.style = "Table Grid"
        headers = ["Total Post", "Sudah Diproses", "Near-Duplicate"]
        values = [
            str(data.total_posts),
            str(data.processed_posts),
            str(data.near_duplicates),
        ]
        for i, h in enumerate(headers):
            cell = stat_table.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "1e3a5f")
        for i, v in enumerate(values):
            cell = stat_table.cell(1, i)
            cell.text = v
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(16)
            cell.paragraphs[0].runs[0].font.color.rgb = COLOR_ACCENT
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if data.language_breakdown:
            lang_p = doc.add_paragraph()
            lang_p.add_run("Bahasa: ").bold = True
            lang_p.add_run("  |  ".join(
                f"{lang.upper()}: {cnt}" for lang, cnt in data.language_breakdown.items()
            ))

        # ── Analisis Sentimen ───────────────────────────────────────────────────
        doc.add_heading("Analisis Sentimen", level=1)

        if data.sentiment.total_analyzed == 0:
            doc.add_paragraph("Belum ada data sentimen.")
        else:
            sent_p = doc.add_paragraph()
            sent_p.add_run(
                f"Total dianalisis: {data.sentiment.total_analyzed} post. "
                f"Dominan: "
            )
            dom_run = sent_p.add_run(data.sentiment.dominant.upper())
            dom_run.bold = True
            dom_run.font.color.rgb = COLOR_PRIMARY

            # Tabel distribusi
            dist_table = doc.add_table(rows=1 + len(data.sentiment.distribution), cols=3)
            dist_table.style = "Table Grid"
            headers = ["Label Sentimen", "Jumlah", "Persentase"]
            for i, h in enumerate(headers):
                cell = dist_table.cell(0, i)
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                _set_cell_bg(cell, "1e3a5f")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            label_color_map = {
                "positive": "27ae60",
                "negative": "e74c3c",
                "neutral": "95a5a6",
            }
            for row_idx, label in enumerate(["positive", "negative", "neutral"], 1):
                cnt = data.sentiment.distribution.get(label, 0)
                pct = data.sentiment.percentages.get(label, 0.0)
                row = dist_table.rows[row_idx]
                row.cells[0].text = label.capitalize()
                row.cells[1].text = str(cnt)
                row.cells[2].text = f"{pct:.1f}%"
                hex_c = label_color_map.get(label, "999999")
                for cell in row.cells:
                    _set_cell_bg(cell, "f8f9fa")

            # Contoh post
            if data.sentiment.examples:
                doc.add_paragraph()
                doc.add_paragraph("Contoh post per sentimen:").runs[0].bold = True
                for ex in data.sentiment.examples:
                    p = doc.add_paragraph(style="List Bullet")
                    label_run = p.add_run(f"[{ex['label'].upper()}] ")
                    label_run.bold = True
                    color_map_rgb = {
                        "positive": COLOR_POSITIVE,
                        "negative": COLOR_NEGATIVE,
                        "neutral": COLOR_NEUTRAL,
                    }
                    label_run.font.color.rgb = color_map_rgb.get(ex["label"], COLOR_NEUTRAL)
                    snippet = ex["content"][:200]
                    p.add_run(f"{ex['platform']} — {snippet}")

        # ── Entitas Terdeteksi ──────────────────────────────────────────────────
        doc.add_heading("Entitas Terdeteksi", level=1)

        if not data.entities.by_type:
            doc.add_paragraph("Belum ada data entitas.")
        else:
            doc.add_paragraph(f"Total entitas unik: {data.entities.total_unique}")
            for etype, items in data.entities.by_type.items():
                if not items:
                    continue
                doc.add_heading(etype, level=2)
                ent_table = doc.add_table(rows=1 + len(items[:10]), cols=2)
                ent_table.style = "Table Grid"
                ent_table.cell(0, 0).text = "Entitas"
                ent_table.cell(0, 1).text = "Jumlah"
                for cell in ent_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].bold = True
                    _set_cell_bg(cell, "e8ecf0")
                for row_idx, item in enumerate(items[:10], 1):
                    ent_table.cell(row_idx, 0).text = item["text"]
                    ent_table.cell(row_idx, 1).text = str(item["count"])

        # ── Tren Aktivitas ──────────────────────────────────────────────────────
        doc.add_heading("Tren Aktivitas", level=1)

        if not data.trend.volume:
            doc.add_paragraph("Belum ada data tren.")
        else:
            trend_p = doc.add_paragraph()
            trend_p.add_run(f"Arah tren: ")
            dir_run = trend_p.add_run(data.trend.direction.upper())
            dir_run.bold = True
            dir_run.font.color.rgb = COLOR_ACCENT
            trend_p.add_run(f"  |  Total: {data.trend.total_posts} post")

            if data.trend.platform_breakdown:
                p = doc.add_paragraph()
                p.add_run("Platform: ").bold = True
                p.add_run("  |  ".join(
                    f"{plat}: {cnt}" for plat, cnt in data.trend.platform_breakdown.items()
                ))

            # Tabel volume — max 10 baris
            vol_rows = data.trend.volume[:10]
            vol_table = doc.add_table(rows=1 + len(vol_rows), cols=3)
            vol_table.style = "Table Grid"
            for h_idx, h in enumerate(["Periode", "Platform", "Jumlah Post"]):
                cell = vol_table.cell(0, h_idx)
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                _set_cell_bg(cell, "1e3a5f")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for r_idx, v in enumerate(vol_rows, 1):
                vol_table.cell(r_idx, 0).text = v["period"][:10]
                vol_table.cell(r_idx, 1).text = v["platform"]
                vol_table.cell(r_idx, 2).text = str(v["count"])

        # ── Sample Posts ────────────────────────────────────────────────────────
        if data.top_posts:
            doc.add_heading("Contoh Post Terbaru", level=1)
            for i, post in enumerate(data.top_posts, 1):
                label_info = (
                    f" [{post['sentiment_label'].upper()}]" if post.get("sentiment_label") else ""
                )
                h_p = doc.add_paragraph()
                h_run = h_p.add_run(
                    f"{i}. {post.get('platform', '').upper()}{label_info} — "
                    f"{post.get('author', 'unknown')} — {(post.get('published_at') or '')[:10]}"
                )
                h_run.bold = True
                h_run.font.size = Pt(9)
                content_snippet = (post.get("content") or "")[:300]
                q_p = doc.add_paragraph(content_snippet)
                q_p.paragraph_format.left_indent = Cm(0.5)
                for run in q_p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # ── Footer ──────────────────────────────────────────────────────────────
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            "Social Intelligence Platform — Laporan ini dibuat otomatis oleh sistem AI."
        )
        footer_run.font.size = Pt(7)
        footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        doc.save(file_path)
        return file_path
